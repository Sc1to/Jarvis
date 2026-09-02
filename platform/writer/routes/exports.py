import json
import os
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import db
from deps import current_user

router = APIRouter()


def _exports_dir(book_id: str) -> str:
    return os.path.join(db.data_dir(book_id), "exports")

def _index_path(book_id: str) -> str:
    return os.path.join(_exports_dir(book_id), "index.json")

def _export_path(book_id: str, export_id: str) -> str:
    return os.path.join(_exports_dir(book_id), f"{export_id}.json")

def _read_index(book_id: str) -> list:
    p = _index_path(book_id)
    return json.load(open(p)) if os.path.exists(p) else []

def _save_index(book_id: str, index: list) -> None:
    os.makedirs(_exports_dir(book_id), exist_ok=True)
    with open(_index_path(book_id), "w") as f:
        json.dump(index, f, indent=2)

def _read_export(book_id: str, export_id: str) -> dict | None:
    p = _export_path(book_id, export_id)
    return json.load(open(p)) if os.path.exists(p) else None

def _save_export(book_id: str, export: dict) -> None:
    os.makedirs(_exports_dir(book_id), exist_ok=True)
    with open(_export_path(book_id, export["export_id"]), "w") as f:
        json.dump(export, f, indent=2)

def _chapter_path(book_id: str, chapter: int) -> str:
    return os.path.join(db.data_dir(book_id), f"chapter_{chapter:02d}.md")

def _chapter_meta_path(book_id: str, chapter: int) -> str:
    return os.path.join(db.data_dir(book_id), f"chapter_{chapter:02d}_meta.json")

def _get_chapter_title(book_id: str, chapter: int) -> str:
    tier4_status_path = os.path.join(db.data_dir(book_id), "tier4", "status.json")
    if os.path.exists(tier4_status_path):
        try:
            for ch in json.load(open(tier4_status_path)).get("chapters", []):
                if ch["number"] == chapter:
                    return ch.get("title", f"Chapter {chapter}")
        except Exception:
            pass
    return f"Chapter {chapter}"

def _word_count(text: str) -> int:
    return len(text.split())

def _total_words(chapters: list) -> int:
    return sum(_word_count(ch.get("content", "")) for ch in chapters)


@router.post("/books/{book_id}/exports")
def create_export(book_id: str, user: str = Depends(current_user)):
    chapters = []
    i = 1
    while True:
        meta_path = _chapter_meta_path(book_id, i)
        chapter_path = _chapter_path(book_id, i)
        if not os.path.exists(chapter_path) and not os.path.exists(meta_path):
            break
        if os.path.exists(meta_path) and os.path.exists(chapter_path):
            meta = json.load(open(meta_path))
            if meta.get("status") == "approved":
                content = open(chapter_path).read()
                chapters.append({
                    "chapter": i,
                    "title": _get_chapter_title(book_id, i),
                    "content": content,
                    "word_count": _word_count(content),
                })
        i += 1

    if not chapters:
        return {"error": "No approved chapters to export", "status": "error"}

    now = datetime.now(timezone.utc).isoformat()
    export_id = str(uuid.uuid4())[:8]
    label = f"Export {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')}"
    total_words = _total_words(chapters)

    export = {
        "export_id": export_id,
        "book_id": book_id,
        "label": label,
        "created_at": now,
        "updated_at": now,
        "chapter_count": len(chapters),
        "word_count": total_words,
        "chapters": chapters,
    }
    _save_export(book_id, export)

    index = _read_index(book_id)
    index.insert(0, {
        "export_id": export_id,
        "label": label,
        "created_at": now,
        "updated_at": now,
        "chapter_count": len(chapters),
        "word_count": total_words,
    })
    _save_index(book_id, index)

    return {"export_id": export_id, "status": "ok"}


@router.get("/books/{book_id}/exports")
def list_exports(book_id: str):
    return {"exports": _read_index(book_id), "status": "ok"}


@router.get("/books/{book_id}/exports/{export_id}")
def get_export(book_id: str, export_id: str):
    export = _read_export(book_id, export_id)
    if not export:
        return {"error": "Export not found", "status": "error"}
    return {"export": export, "status": "ok"}


class UpdateExportBody(BaseModel):
    label: str | None = None
    chapters: list | None = None


@router.put("/books/{book_id}/exports/{export_id}")
def update_export(book_id: str, export_id: str, body: UpdateExportBody, user: str = Depends(current_user)):
    export = _read_export(book_id, export_id)
    if not export:
        return {"error": "Export not found", "status": "error"}

    now = datetime.now(timezone.utc).isoformat()
    if body.label is not None:
        export["label"] = body.label
    if body.chapters is not None:
        for ch in body.chapters:
            ch["word_count"] = _word_count(ch.get("content", ""))
        export["chapters"] = body.chapters
        export["chapter_count"] = len(body.chapters)
        export["word_count"] = _total_words(body.chapters)
    export["updated_at"] = now
    _save_export(book_id, export)

    index = _read_index(book_id)
    for entry in index:
        if entry["export_id"] == export_id:
            entry["label"] = export["label"]
            entry["updated_at"] = now
            entry["word_count"] = export["word_count"]
            entry["chapter_count"] = export["chapter_count"]
            break
    _save_index(book_id, index)

    return {"status": "ok"}


@router.delete("/books/{book_id}/exports/{export_id}")
def delete_export(book_id: str, export_id: str, user: str = Depends(current_user)):
    p = _export_path(book_id, export_id)
    if os.path.exists(p):
        os.remove(p)
    _save_index(book_id, [e for e in _read_index(book_id) if e["export_id"] != export_id])
    return {"status": "ok"}


@router.get("/books/{book_id}/exports/{export_id}/download")
def download_export(book_id: str, export_id: str, format: str = "md"):
    export = _read_export(book_id, export_id)
    if not export:
        return {"error": "Export not found", "status": "error"}

    chapters = export.get("chapters", [])
    safe_label = re.sub(r'[^\w\s-]', '', export["label"]).strip().replace(' ', '_')

    if format == "md":
        parts = [f"# {ch['title']}\n\n{ch['content']}" for ch in chapters]
        content = "\n\n---\n\n".join(parts)
        return StreamingResponse(
            iter([content.encode("utf-8")]),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{safe_label}.md"'},
        )

    if format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {"error": "python-docx not installed", "status": "error"}

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        for i, ch in enumerate(chapters):
            if i > 0:
                doc.add_page_break()
            heading = doc.add_heading(ch["title"], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for para_text in ch["content"].split("\n\n"):
                para_text = para_text.strip()
                if not para_text or para_text == "---" or para_text.startswith("## "):
                    continue
                p = doc.add_paragraph(para_text)
                p.style.font.size = Pt(12)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            iter([buf.read()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_label}.docx"'},
        )

    return {"error": f"Unknown format: {format}", "status": "error"}
